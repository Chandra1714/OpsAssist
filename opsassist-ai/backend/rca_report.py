from io import BytesIO
from datetime import datetime

from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas


def build_rca_docx(incident, analysis, resolution_notes: str) -> bytes:
    document = Document()
    document.add_heading(f"RCA Report - {incident.incident_uid}", 0)
    document.add_paragraph(f"Application: {incident.application_name or 'N/A'}")
    document.add_paragraph(f"Environment: {incident.namespace or 'N/A'}")
    document.add_paragraph(f"Issue Summary: {incident.description}")
    document.add_heading("Root Cause", level=1)
    document.add_paragraph(analysis.get("root_cause", "Unknown"))
    document.add_heading("Evidence", level=1)
    document.add_paragraph(incident.error_message or "No error message.")
    if incident.logs:
        document.add_paragraph("Logs:")
        document.add_paragraph(incident.logs[:3000])
    document.add_heading("Resolution", level=1)
    document.add_paragraph(resolution_notes or "No resolution notes provided.")
    document.add_heading("Duration", level=1)
    document.add_paragraph(analysis.get("duration", "Unknown"))
    document.add_heading("Preventive Actions", level=1)
    for item in analysis.get("preventive_actions", ["Review incident process."]):
        document.add_paragraph(f"- {item}")
    bio = BytesIO()
    document.save(bio)
    return bio.getvalue()


def build_rca_pdf(incident, analysis, resolution_notes: str) -> bytes:
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    margin = 40
    y = height - margin
    c.setFont("Helvetica-Bold", 20)
    c.drawString(margin, y, f"RCA Report - {incident.incident_uid}")
    y -= 40
    c.setFont("Helvetica", 11)
    lines = [
        f"Application: {incident.application_name or 'N/A'}",
        f"Environment: {incident.namespace or 'N/A'}",
        f"Issue Summary: {incident.description}",
        "",
        "Root Cause:",
        analysis.get("root_cause", "Unknown"),
        "",
        "Evidence:",
        incident.error_message or "No error message.",
        "",
        "Resolution:",
        resolution_notes or "No resolution notes provided.",
        "",
        "Duration:",
        analysis.get("duration", "Unknown"),
        "",
        "Preventive Actions:",
    ]
    for item in analysis.get("preventive_actions", ["Review incident process."]):
        lines.append(f"- {item}")
    for line in lines:
        if y < 60:
            c.showPage()
            y = height - margin
            c.setFont("Helvetica", 11)
        c.drawString(margin, y, line)
        y -= 16
    c.save()
    return buffer.getvalue()
